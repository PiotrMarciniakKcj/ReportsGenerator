from datetime import datetime, timedelta
from calendar import monthcalendar
from docx import Document
from docx.shared import Cm, Pt
from openpyxl import load_workbook
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import csv
import codecs
import urllib.request
import urllib.error
import sys


# convert float to string value with '%'
def float_to_percent(value):
    return ('{:.2f}'.format(value * 100) + "%").replace('.', ',')


# return index of the last row of the table
def enumerate_sheet(worksheet, max_val):
    for i, i_row in enumerate(worksheet):
        if i_row[max_val].value == "Suma:":
            return str(i + 1)
        if i_row[max_val].value == "Suma":
            return str(i + 10)


# return formatted tables from an Excel file
def get_formatted_data(path):
    wb = load_workbook(path)
    sheet = wb.worksheets[0]
    tables_from_excel = [format_data("I2:N" + enumerate_sheet(sheet, 9), path),
                         format_data("R2:X" + enumerate_sheet(sheet, 18), path)]
    return tables_from_excel


def get_formatted_classification_data(path):
    wb = load_workbook(path)
    sheet = wb.worksheets[1]
    tables_from_excel = [format_data("M1:T" + enumerate_sheet(sheet, 13), path, True),
                         format_data("V1:AC" + enumerate_sheet(sheet, 22), path, True)]
    return tables_from_excel


def get_formatted_classification_summary_data(path):
    tables_from_excel = [format_data("C35:H43", path, True), format_data("L35:Q42", path, True)]
    return tables_from_excel


# return formatted table(without headers) from Excel as a 2-dimensional array from an index input
def format_data(indexes, path, classification=False):
    wb = load_workbook(path)
    if classification is False:
        sheet = wb.worksheets[0]
    else:
        sheet = wb.worksheets[1]
    input_table = sheet[indexes]
    new_table = []
    for row in input_table:
        temp_row = []
        if classification is False and row[0].value == row[-1].value:
            continue
        for cell in row:
            cell = cell.value
            if cell is None:
                cell = ''
            if type(cell) is datetime.time:
                cell = cell.strftime("%H:%M:%S")
            temp_row.append(cell)
        row = temp_row
        if row[-1] not in ("-", "%", ""):
            row[-1] = float_to_percent(row[-1])
        new_table.append(row)
    return new_table


# get YYYYMMDD date from the name of the Excel file
def get_date(filename):
    date = filename[filename.find("_") + 1:filename.rfind("_", 0, filename.rfind("_"))]
    date = datetime(int(date[0:4]), int(date[4:6]), int(date[6:8]), int(date[9:11]), int(date[11:13]),
                    int(date[13:15]))
    return date


def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)


def paste_classification_tables(formatted_data, table_name):
    for i, data_row in enumerate(formatted_data):
        table_row = table_name.add_row()
        table_row.height = Cm(0.5)
        table_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for x in range(0, len(data_row)):
            table_row.cells[x].text = str(data_row[x])
            table_row.cells[x].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            table_row.cells[x].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraphs = table_row.cells[x].paragraphs
            paragraph = paragraphs[0]
            run_obj = paragraph.runs
            run = run_obj[0]
            font = run.font
            font.size = Pt(8)


        for y in range(0, len(table_name.rows) // 11):
            for z in range(0, 2):
                a = table_name.rows[y * 11 + 1].cells[z]
                b = table_name.rows[y * 11 + 10].cells[z]
                a.merge(b)
                a.text = a.text.strip()
                a.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    remove_row(table_name, table_name.rows[0])
    widths = (Cm(2), Cm(2), Cm(5), Cm(1.3), Cm(1.3), Cm(1.3), Cm(1.3), Cm(1.3))
    for row in table_name.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
    for idx, col in enumerate(table_name.columns):
        col.width = widths[idx]
    table_name.alignment = WD_TABLE_ALIGNMENT.CENTER


def paste_cost323_classification_tables(formatted_data, table_name):
    for i, data_row in enumerate(formatted_data):
        table_row = table_name.add_row()
        table_row.height = Cm(0.5)
        table_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for x in range(0, len(data_row)):
            table_row.cells[x].text = str(data_row[x])
            table_row.cells[x].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            table_row.cells[x].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraphs = table_row.cells[x].paragraphs
            paragraph = paragraphs[0]
            run_obj = paragraph.runs
            run = run_obj[0]
            font = run.font
            font.size = Pt(8)

        for y in range(0, len(table_name.rows) // 11):
            for z in range(0, 2):
                a = table_name.rows[y * 11 + 1].cells[z]
                b = table_name.rows[y * 11 + 9].cells[z]
                a.merge(b)
                a.text = a.text.strip()
                a.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    remove_row(table_name, table_name.rows[0])
    widths = (Cm(2), Cm(2), Cm(5), Cm(1.3), Cm(1.3), Cm(1.3), Cm(1.3), Cm(1.3))
    for row in table_name.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
    for idx, col in enumerate(table_name.columns):
        col.width = widths[idx]
    remove_row(table_name, table_name.rows[-1])
    table_name.alignment = WD_TABLE_ALIGNMENT.CENTER


# copy the contents of the tables from Excel and paste them into the word template
def paste_summary_table(formatted_data, table_name):
    for i, data_row in enumerate(formatted_data):
        for x in range(0, 2):
            table_name.rows[x].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            table_name.rows[x].height = Cm(1)
            for y in range(0, 4):
                table_name.rows[x].cells[y].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row = table_name.rows[1].cells
        for x in range(0, 4):
            row[x].text = str(data_row[x])
            row[x].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            row[x].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    table_name.alignment = WD_TABLE_ALIGNMENT.CENTER



# naglowki sie usuwaja
def paste_classification_summary_table(formatted_data, table_name):
    for i, data_row in enumerate(formatted_data):
        table_row = table_name.add_row()
        table_row.height = Cm(0.7)
        table_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for x in range(len(data_row)):
            table_row.cells[x].text = str(data_row[x])
            table_row.cells[x].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            table_row.cells[x].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table_name.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = (Cm(6), Cm(2), Cm(2), Cm(2), Cm(2), Cm(2))
    for row in table_name.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


# copy the contents of the tables from Excel and paste them into the word template
def paste_table(formatted_data, table_name):
    for i, data_row in enumerate(formatted_data):
        table_row = table_name.add_row()
        table_row.height = Cm(0.5)
        table_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for x in range(0, len(data_row)):
            table_row.cells[x].text = str(data_row[x])
            table_row.cells[x].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            table_row.cells[x].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table_name.alignment = WD_TABLE_ALIGNMENT.CENTER


# get weather
def get_weather(date, x_coordinate, y_coordinate):
    # This is the core of our weather query URL
    BaseURL = 'https://weather.visualcrossing.com/VisualCrossingWebServices/rest/services/timeline/'

    ApiKey = '5ZRVAYDXFUCTLPHARB5M9JMVJ'
    # UnitGroup sets the units of the output - us or metric
    UnitGroup = 'metric'

    # Location for the weather data
    Location = str(x_coordinate) + "," + str(y_coordinate)

    date = date + timedelta(hours=2)
    # Optional start and end dates
    # If nothing is specified, the forecast is retrieved.
    # If start date only is specified, a single historical or forecast day will be retrieved
    # If both start and end date are specified, a date range will be retrieved
    StartDate = date.strftime("%Y-%m-%dT%H:%M:%S")
    EndDate = ''

    # JSON or CSV
    # JSON format supports daily, hourly, current conditions, weather alerts and events in a single JSON package
    # CSV format requires an 'include' parameter below to indicate which table section is required
    ContentType = "csv"

    # include sections
    # values include days,hours,current,alerts
    Include = "current&elements=datetime,temp,conditions"

    # basic query including location
    ApiQuery = BaseURL + Location

    # append the start and end date if present
    if len(StartDate):
        ApiQuery += "/" + StartDate
        if len(EndDate):
            ApiQuery += "/" + EndDate

    # Url is completed. Now add query parameters (could be passed as GET or POST)
    ApiQuery += "?"

    # append each parameter as necessary
    if len(UnitGroup):
        ApiQuery += "&unitGroup=" + UnitGroup

    if len(ContentType):
        ApiQuery += "&contentType=" + ContentType

    if len(Include):
        ApiQuery += "&include=" + Include
    ApiQuery += "&lang=pl"
    ApiQuery += "&key=" + ApiKey

    # print(' - Running query URL: ', ApiQuery)
    # print()

    try:
        CSVBytes = urllib.request.urlopen(ApiQuery)
    except urllib.error.HTTPError as e:
        ErrorInfo = e.read().decode()
        print('Error code: ', e.code, ErrorInfo)
        sys.exit()
    except urllib.error.URLError as e:
        ErrorInfo = e.read().decode()
        print('Error code: ', e.code, ErrorInfo)
        sys.exit()

    # Parse the results as CSV
    CSVText = csv.reader(codecs.iterdecode(CSVBytes, 'utf-8'))

    RowIndex = 0

    # The first row contain the headers and the additional rows each contain the weather metrics for a single day
    for Row in CSVText:
        if RowIndex == 0:
            FirstRow = Row
        else:
            # print('Weather in ', Row[0], ' on ', Row[1], ', ', Row[2])
            return [Row[1], Row[2]]
            ColIndex = 0
            for Col in Row:
                if ColIndex >= 4:
                    print('   ', FirstRow[ColIndex], ' = ', Row[ColIndex])
                ColIndex += 1
        RowIndex += 1

    # If there are no CSV rows then something fundamental went wrong
    if RowIndex == 0:
        print('Sorry, but it appears that there was an error connecting to the weather server.')
        print('Please check your network connection and try again..')

    # If there is only one CSV  row then we likely got an error from the server
    if RowIndex == 1:
        print('Sorry, but it appears that there was an error retrieving the weather data.')
        print('Error: ', FirstRow)

    print()


# get last sunday of the month
def get_last_sunday(date, input_month):
    month = monthcalendar(date.year, input_month)
    if month[-1][6]:
        return month[-1][-2]
    else:
        return month[-2][-1]


# check whether it is the winter time or summer time and return how many hours to add
def add_hours(date):
    winter_time_start = datetime(date.year, 10, get_last_sunday(date, 10), 0)
    winter_time_end = datetime(date.year, 3, get_last_sunday(date, 3), 0)
    if date > winter_time_start or date < winter_time_end:
        return 1
    return 2


# return a formatted string of the test start description
def get_test_start_text(date, point_x, point_y):
    weather = get_weather(date, point_x, point_y)
    text = "Rozpoczęcie testu: " + date.strftime("%d.%m.%Y") + ", godz. " + (
            date + timedelta(hours=add_hours(date))).strftime("%H:%M:%S") + " (" + date.strftime(
        "%H:%M:%S") + " UTC); temp: ok. " + str(round(float(weather[0]))) + " st., " + weather[1] + "."
    return text


# replace template elements with correct ones
def format_document(document, DR_date, DP_date, N_date, order, summary, summary2=None):
    DR_text = 'Badanie w okresie przed południem (DR500)'
    DP_text = 'Badanie w okresie po południu (DP500)'
    N_text = 'Badanie w okresie nocnym (N200)'

    x = 52.07199
    y = 17.23126
    DR_test_start_text = get_test_start_text(DR_date, x, y)
    DP_test_start_text = get_test_start_text(DP_date, x, y)
    N_test_start_text = get_test_start_text(N_date, x, y)

    if order[0] == DR_date:
        text_list = [DR_text, DP_text, N_text]
        test_start_text_list = [DR_test_start_text, DP_test_start_text, N_test_start_text]
    elif order[0] == DP_date:
        text_list = [DP_text, N_text, DR_text]
        test_start_text_list = [DP_test_start_text, N_test_start_text, DR_test_start_text]
    else:
        text_list = [N_text, DR_text, DP_text]
        test_start_text_list = [N_test_start_text, DR_test_start_text, DP_test_start_text]
    for paragraph in document.paragraphs:
        if 'order1' in paragraph.text:
            paragraph.text = text_list[0]
        if 'order2' in paragraph.text:
            paragraph.text = text_list[1]
        if 'order3' in paragraph.text:
            paragraph.text = text_list[2]
        if 'description1' in paragraph.text:
            paragraph.text = test_start_text_list[0]
        if 'description2' in paragraph.text:
            paragraph.text = test_start_text_list[1]
        if 'description3' in paragraph.text:
            paragraph.text = test_start_text_list[2]
        if 'today' in paragraph.text:
            paragraph.text = "Poznań, " + datetime.today().strftime("%d.%m.%Y")
        if 'summary_d' in paragraph.text:
            paragraph.text = check_d_and_r(summary)[0]
        if 'summary_r' in paragraph.text:
            paragraph.text = check_d_and_r(summary)[1]
        if summary2 is not None:
            if 'summary_d2' in paragraph.text:
                paragraph.text = check_d_and_r(summary2)[0]
            if 'summary_r2' in paragraph.text:
                paragraph.text = check_d_and_r(summary2)[1]

    body_elements = document._body._body
    rs = body_elements.xpath('//w:r')
    # changing elements in the table of contents
    for r in rs:
        if r.text == "toc1":
            r.text = text_list[0]
        if r.text == "toc2":
            r.text = text_list[1]
        if r.text == "toc3":
            r.text = text_list[2]


# get summary tables
def get_summary_tables(tables):
    N_sum = 0
    em_sum = 0
    ef_sum = 0

    Nid_sum = 0
    Kok_sum = 0
    rejected_sum = 0
    for table in tables:
        if tables.index(table) % 2 == 0:
            N_sum += table[-1][2]
            em_sum += table[-1][3]
            ef_sum += table[-1][4]
        else:
            Nid_sum += table[-1][2]
            Kok_sum += table[-1][3]
            rejected_sum += table[-1][5]

    d = float_to_percent((N_sum - em_sum - ef_sum) / N_sum)
    r = float_to_percent(Kok_sum / Nid_sum)

    # 2-dimensional array with just one row
    summary_detection = [[N_sum, em_sum, ef_sum, d]]
    summary_identification = [[Nid_sum, Kok_sum, r, rejected_sum]]
    return [summary_detection, summary_identification]


# check and write to file whether OPZ requirements are met
def check_d_and_r(summary):
    check_d = True
    check_r = True
    if float(summary[0][0][3].replace(",", ".").replace("%", "")) < 99:
        check_d = False
    if float(summary[1][0][2].replace(",", ".").replace("%", "")) < 90:
        check_r = False
    if check_d:
        summary_d_text = "System spełnia wymagania OPZ w zakresie poziomów detekcji."
    else:
        summary_d_text = "System nie spełnia wymagań OPZ w zakresie poziomów detekcji."
    if check_r:
        summary_r_text = "System spełnia wymagania OPZ w zakresie poziomów identyfikacji tablic rejestracyjnych."
    else:
        summary_r_text = "System nie spełnia wymagań OPZ w zakresie poziomów identyfikacji tablic rejestracyjnych."
    return summary_d_text, summary_r_text


# generate report with detection and identification only
def generate_detection_and_identification_report(DR_excel, DP_excel, N_excel):
    # dates of the tests
    DR_date = get_date(DR_excel)
    DP_date = get_date(DP_excel)
    N_date = get_date(N_excel)

    # correct order of the tests
    order = [DR_date, DP_date, N_date]
    order.sort()

    # path to template file
    template = "Raport z testu kamer ANPR Krzykosy_template.docx"
    output = template[0:-14] + " " + datetime.today().strftime("%Y") + ".docx"

    # getting the tables from the Excel files
    formatted_tables = get_formatted_data(DR_excel)
    formatted_tables.extend(get_formatted_data(DP_excel))
    formatted_tables.extend(get_formatted_data(N_excel))
    summary_tables = get_summary_tables(formatted_tables)

    doc = Document(template)
    tables = doc.tables[2:]

    # assigning correct table order
    DR_detection_table = tables[(order.index(DR_date) * 2)]
    DR_identification_table = tables[(order.index(DR_date) * 2 + 1)]
    DP_detection_table = tables[(order.index(DP_date) * 2)]
    DP_identification_table = tables[(order.index(DP_date) * 2 + 1)]
    N_detection_table = tables[(order.index(N_date) * 2)]
    N_identification_table = tables[(order.index(N_date) * 2 + 1)]
    summary_detection_table = tables[-2]
    summary_identification_table = tables[-1]

    # pasting the tables
    paste_table(formatted_tables[0], DR_detection_table)
    paste_table(formatted_tables[1], DR_identification_table)
    paste_table(formatted_tables[2], DP_detection_table)
    paste_table(formatted_tables[3], DP_identification_table)
    paste_table(formatted_tables[4], N_detection_table)
    paste_table(formatted_tables[5], N_identification_table)
    paste_summary_table(summary_tables[0], summary_detection_table)
    paste_summary_table(summary_tables[1], summary_identification_table)

    format_document(doc, DR_date, DP_date, N_date, order, summary_tables)
    doc.save(output)


# generate report with classification
def generate_classification_report(paths):
    # dates of the tests
    DR_date = get_date(paths[0])
    DP_date = get_date(paths[1])
    N_date = get_date(paths[2])

    # DR_date2 = get_date(paths[3])
    # DP_date2 = get_date(paths[4])
    # N_date2 = get_date(paths[5])

    # correct order of the tests
    order = [DR_date, DP_date, N_date]
    order.sort()

    # path to template file
    template = "Raport z testu kamer ANPR Mostki_template.docx"
    output = template[0:-14] + " " + datetime.today().strftime("%Y") + ".docx"

    # getting the tables from the Excel files
    formatted_tables = get_formatted_data(paths[0])
    formatted_tables.extend(get_formatted_data(paths[1]))
    formatted_tables.extend(get_formatted_data(paths[2]))
    classification_tables = get_formatted_classification_data(paths[0])
    classification_tables.extend(get_formatted_classification_data(paths[1]))
    classification_tables.extend(get_formatted_classification_data(paths[2]))
    summary_tables = get_summary_tables(formatted_tables)
    formatted_tables2 = get_formatted_data(paths[3])
    formatted_tables2.extend(get_formatted_data(paths[4]))
    formatted_tables2.extend(get_formatted_data(paths[5]))
    classification_tables2 = get_formatted_classification_data(paths[3])
    classification_tables2.extend(get_formatted_classification_data(paths[4]))
    classification_tables2.extend(get_formatted_classification_data(paths[5]))
    summary_tables2 = get_summary_tables(formatted_tables2)
    classification_summary_tables = get_formatted_classification_summary_data(paths[6])
    classification_summary_tables2 = get_formatted_classification_summary_data(paths[7])

    doc = Document(template)
    tables = doc.tables[2:]

    # getting the identification tables in word
    DR_identification_table = tables[order.index(DR_date)]
    DP_identification_table = tables[order.index(DP_date)]
    N_identification_table = tables[order.index(N_date)]
    DR_identification_table2 = tables[order.index(DR_date) + 4]
    DP_identification_table2 = tables[order.index(DP_date) + 4]
    N_identification_table2 = tables[order.index(N_date) + 4]
    summary_identification_table = tables[3]
    summary_identification_table2 = tables[7]

    # pasting the identification tables
    paste_table(formatted_tables[1], DR_identification_table)
    paste_table(formatted_tables[3], DP_identification_table)
    paste_table(formatted_tables[5], N_identification_table)
    paste_summary_table(summary_tables[1], summary_identification_table)

    paste_table(formatted_tables2[1], DR_identification_table2)
    paste_table(formatted_tables2[3], DP_identification_table2)
    paste_table(formatted_tables2[5], N_identification_table2)
    paste_summary_table(summary_tables2[1], summary_identification_table2)

    # getting the detection tables in word
    DR_detection_table = tables[order.index(DR_date) * 3 + 10]
    DP_detection_table = tables[order.index(DP_date) * 3 + 10]
    N_detection_table = tables[order.index(N_date) * 3 + 10]
    summary_detection_table = tables[19]

    DR_detection_table2 = tables[order.index(DR_date) * 3 + 22]
    DP_detection_table2 = tables[order.index(DP_date) * 3 + 22]
    N_detection_table2 = tables[order.index(N_date) * 3 + 22]
    summary_detection_table2 = tables[31]

    # pasting the classification tables in word
    for i in range(3):
        paste_classification_tables(classification_tables[i * 2], tables[3 * i + 8])
        paste_cost323_classification_tables(classification_tables[i * 2 + 1], tables[3 * i + 9])

        paste_classification_tables(classification_tables2[i * 2], tables[3 * i + 20])
        paste_cost323_classification_tables(classification_tables2[i * 2 + 1], tables[3 * i + 21])

    paste_classification_summary_table(classification_summary_tables[0], tables[17])
    paste_classification_summary_table(classification_summary_tables[1], tables[18])
    #paste_cost323_classification_summary_table(classification_summary_tables[1], tables[18])
    paste_classification_summary_table(classification_summary_tables2[0], tables[29])
    paste_classification_summary_table(classification_summary_tables2[1], tables[30])
    #paste_cost323_classification_summary_table(classification_summary_tables2[1], tables[30])

    # pasting the detection tables in word
    paste_table(formatted_tables[0], DR_detection_table)
    paste_table(formatted_tables[2], DP_detection_table)
    paste_table(formatted_tables[4], N_detection_table)
    paste_summary_table(summary_tables[0], summary_detection_table)

    paste_table(formatted_tables2[0], DR_detection_table2)
    paste_table(formatted_tables2[2], DP_detection_table2)
    paste_table(formatted_tables2[4], N_detection_table2)
    paste_summary_table(summary_tables2[0], summary_detection_table2)

    format_document(doc, DR_date, DP_date, N_date, order, summary_tables, summary_tables2)
    doc.save(output)


# DR_path = "102_20230516_090002_DR500_wClass.xlsx"
# DP_path = "102_20230516_140007_DP500_wClass.xlsx"
# N_path = "102_20230515_210000_N200_wClass.xlsx"
# generate_detection_and_identification_report(DR_excel=DR_path, DP_excel=DP_path, N_excel=N_path)

DR_path = "54_20230524_050101_DR500_wClass.xlsx"
DP_path = "54_20230523_123001_DP500_wClass.xlsx"
N_path = "54_20230523_205800_N200_wClass.xlsx"
DR_path2 = "55_20230524_050000_DR500_wClass.xlsx"
DP_path2 = "55_20230523_123000_DP500_wClass.xlsx"
N_path2 = "55_20230523_205804_N200_wClass.xlsx"
summary1 = "54.xlsx"
summary2 = "55.xlsx"
generate_classification_report([DR_path, DP_path, N_path, DR_path2, DP_path2, N_path2, summary1, summary2])
